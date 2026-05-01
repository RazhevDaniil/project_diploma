"""Document parser — extracts text and tables from PDF, DOCX, XLSX, TXT."""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path

logger = logging.getLogger(__name__)


@dataclass
class ParsedTable:
    """A single table extracted from a document."""
    page_or_section: str
    headers: list[str]
    rows: list[list[str]]

    def to_markdown(self) -> str:
        if not self.headers and not self.rows:
            return ""
        cols = self.headers if self.headers else [f"Col{i}" for i in range(len(self.rows[0]))]
        lines = ["| " + " | ".join(cols) + " |"]
        lines.append("| " + " | ".join(["---"] * len(cols)) + " |")
        for row in self.rows:
            # Pad row to match column count
            padded = row + [""] * (len(cols) - len(row))
            lines.append("| " + " | ".join(padded[:len(cols)]) + " |")
        return "\n".join(lines)


@dataclass
class ParsedBlock:
    """A structured document block in original order."""

    kind: str
    text: str
    style: str = ""
    level: int | None = None
    table_index: int | None = None
    row_index: int | None = None
    headers: list[str] = field(default_factory=list)
    cells: list[str] = field(default_factory=list)


@dataclass
class ParsedDocument:
    """Full result of parsing a document."""
    filename: str
    text: str
    tables: list[ParsedTable] = field(default_factory=list)
    metadata: dict = field(default_factory=dict)
    blocks: list[ParsedBlock] = field(default_factory=list)

    @property
    def full_text(self) -> str:
        """Text with tables appended as markdown."""
        if self.metadata.get("tables_in_text"):
            return self.text
        parts = [self.text]
        for t in self.tables:
            md = t.to_markdown()
            if md:
                parts.append(f"\n[Таблица из {t.page_or_section}]\n{md}")
        return "\n".join(parts)


def parse_document(file_path: str | Path) -> ParsedDocument:
    """Detect format and parse the document."""
    path = Path(file_path)
    suffix = path.suffix.lower()
    parsers = {
        ".pdf": _parse_pdf,
        ".docx": _parse_docx,
        ".doc": _parse_docx,
        ".xlsx": _parse_xlsx,
        ".xls": _parse_xlsx,
        ".txt": _parse_txt,
    }
    parser_fn = parsers.get(suffix)
    if parser_fn is None:
        raise ValueError(f"Unsupported file format: {suffix}")
    logger.info("Parsing %s with %s parser", path.name, suffix)
    return parser_fn(path)


def _parse_pdf(path: Path) -> ParsedDocument:
    import pdfplumber

    text_parts: list[str] = []
    tables: list[ParsedTable] = []

    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)

            for j, raw_table in enumerate(page.extract_tables()):
                if not raw_table or len(raw_table) < 2:
                    continue
                headers = [str(c) if c else "" for c in raw_table[0]]
                rows = []
                for row in raw_table[1:]:
                    rows.append([str(c) if c else "" for c in row])
                tables.append(ParsedTable(
                    page_or_section=f"стр. {i + 1}, таблица {j + 1}",
                    headers=headers,
                    rows=rows,
                ))

    return ParsedDocument(
        filename=path.name,
        text="\n\n".join(text_parts),
        tables=tables,
    )


def _parse_docx(path: Path) -> ParsedDocument:
    from docx import Document
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(str(path))
    text_parts: list[str] = []
    tables: list[ParsedTable] = []
    blocks: list[ParsedBlock] = []
    table_index = 0

    def iter_block_items(parent):
        for child in parent.element.body.iterchildren():
            if child.tag.endswith("}p"):
                yield Paragraph(child, parent)
            elif child.tag.endswith("}tbl"):
                yield Table(child, parent)

    def paragraph_level(paragraph: Paragraph) -> int | None:
        style_name = paragraph.style.name if paragraph.style else ""
        if style_name.lower().startswith("heading"):
            parts = style_name.split()
            if parts and parts[-1].isdigit():
                return int(parts[-1])
        return None

    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            text = " ".join(item.text.split())
            if not text:
                continue
            style_name = item.style.name if item.style else ""
            level = paragraph_level(item)
            blocks.append(ParsedBlock(
                kind="paragraph",
                text=text,
                style=style_name,
                level=level,
            ))
            text_parts.append(text)
            continue

        table_index += 1
        raw_rows = []
        for row in item.rows:
            raw_rows.append([" ".join(cell.text.split()) for cell in row.cells])
        if len(raw_rows) < 2:
            continue
        headers = raw_rows[0]
        rows = raw_rows[1:]
        table = ParsedTable(
            page_or_section=f"таблица {table_index}",
            headers=headers,
            rows=rows,
        )
        tables.append(table)
        text_parts.append(f"[Таблица {table_index}]")
        text_parts.append(table.to_markdown())
        for row_index, row in enumerate(rows, start=1):
            pairs = []
            for idx, cell in enumerate(row):
                header = headers[idx] if idx < len(headers) and headers[idx] else f"Колонка {idx + 1}"
                if cell:
                    pairs.append(f"{header}: {cell}")
            row_text = "; ".join(pairs)
            if row_text:
                blocks.append(ParsedBlock(
                    kind="table_row",
                    text=row_text,
                    table_index=table_index,
                    row_index=row_index,
                    headers=headers,
                    cells=row,
                ))

    return ParsedDocument(
        filename=path.name,
        text="\n".join(text_parts),
        tables=tables,
        metadata={"tables_in_text": True},
        blocks=blocks,
    )


def _parse_xlsx(path: Path) -> ParsedDocument:
    import openpyxl

    wb = openpyxl.load_workbook(str(path), data_only=True)
    text_parts: list[str] = []
    tables: list[ParsedTable] = []

    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        rows_data = []
        for row in ws.iter_rows(values_only=True):
            str_row = [str(c) if c is not None else "" for c in row]
            if any(c.strip() for c in str_row):
                rows_data.append(str_row)
        if not rows_data:
            continue
        headers = rows_data[0]
        rows = rows_data[1:]
        tables.append(ParsedTable(
            page_or_section=f"лист «{sheet_name}»",
            headers=headers,
            rows=rows,
        ))
        # Also add as text
        for row in rows_data:
            text_parts.append(" | ".join(row))

    return ParsedDocument(
        filename=path.name,
        text="\n".join(text_parts),
        tables=tables,
    )


def _parse_txt(path: Path) -> ParsedDocument:
    text = path.read_text(encoding="utf-8", errors="replace")
    return ParsedDocument(filename=path.name, text=text)
