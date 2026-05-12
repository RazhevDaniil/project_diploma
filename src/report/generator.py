"""Report generator — produces Markdown, PDF, and DOCX reports."""

from __future__ import annotations

import logging
import re
from datetime import datetime
from pathlib import Path

from src.models import AnalysisReport, RequirementVerdict

import config as cfg

logger = logging.getLogger(__name__)

VERDICT_LABELS = {
    "match": "Соответствует",
    "partial": "Частично соответствует",
    "mismatch": "Не соответствует",
    "needs_clarification": "Требует уточнения",
    "out_of_scope": "Вне технической оценки",
}

VERDICT_ICONS = {
    "match": "✅",
    "partial": "🟡",
    "mismatch": "❌",
    "needs_clarification": "❓",
    "out_of_scope": "⚪",
}

CATEGORY_LABELS = {
    "technical": "Техническое",
    "sla": "SLA",
    "legal": "Юридическое",
    "commercial": "Коммерческое",
    "security": "Информационная безопасность",
    "procedural": "Процедурное (закупка)",
    "other": "Прочее",
}

KEY_MATCHES_LIMIT = 10
CATEGORY_PRIORITY = {
    "technical": 0,
    "sla": 1,
    "security": 2,
    "legal": 3,
    "commercial": 4,
    "other": 5,
}

PLATFORM_VERDICT_SYMBOLS = {
    "match": "+",
    "partial": "±",
    "mismatch": "-",
    "needs_clarification": "?",
}
PREFERRED_PLATFORM_ORDER = ["ГосОблако", "Evolution", "Advanced", "Облако VMware"]
PLATFORM_VERDICT_RANK = {
    "match": 3,
    "partial": 2,
    "needs_clarification": 1,
    "mismatch": 0,
}


def _section_label(v: RequirementVerdict) -> str:
    """Format the section/point label with fallback."""
    section = v.section.strip() if v.section else ""
    if section:
        return f"Пункт {section}"
    return f"Требование №{v.requirement_id}"


_NUMERIC_UNIT_RE = re.compile(
    r"\d[\d\s.,]*\s*(?:%|гб/?с|тб/?с|мбит/?с|гбит/?с|гб|тб|мб|кб|"
    r"мс|сек|секунд|минут|часов|часа|часов|ч|дней|дня|"
    r"vcpu|cpu|ram|iops|rps|шт|ед|ту|узл|объект|бакет|корзин)\b",
    re.IGNORECASE,
)
_OR_ALTERNATIVE_RE = re.compile(r"\bили\b|\bлибо\b", re.IGNORECASE)


def _req_text(v: RequirementVerdict, max_len: int = 200) -> str:
    """Compact rendering of a requirement preserving numbers and 'или'-clauses.

    Plain `text[:max_len]` cuts off the tail, often losing the actual
    parameter (e.g. «не более 20 Мбит/с», «или репликация в 2 зоны»).
    For matrix rows or report tables we keep the head, but if a critical
    fragment lies beyond max_len, we append a compact suffix
    `head … <numeric or 'или'-fragment>` while staying within max_len.
    """
    text = v.requirement_text.strip() if v.requirement_text else ""
    if not text:
        if v.reasoning:
            return f"[текст не извлечён] {v.reasoning[:max_len]}"
        return "[текст требования не извлечён]"
    if len(text) <= max_len:
        return text

    head_budget = max(60, int(max_len * 0.7))
    head = text[:head_budget].rstrip()
    tail = text[head_budget:]

    # Сохраняем числовые маркеры и «или»-альтернативы, если они в хвосте.
    fragments_to_keep: list[str] = []
    for m in _NUMERIC_UNIT_RE.finditer(tail):
        start = max(0, m.start() - 20)
        end = min(len(tail), m.end() + 10)
        frag = tail[start:end].strip(" .,;:")
        if frag and frag not in fragments_to_keep:
            fragments_to_keep.append(frag)
    or_match = _OR_ALTERNATIVE_RE.search(tail)
    if or_match:
        start = max(0, or_match.start() - 5)
        end = min(len(tail), or_match.end() + 60)
        frag = tail[start:end].strip(" .,;:")
        if frag and frag not in fragments_to_keep:
            fragments_to_keep.append(frag)

    if not fragments_to_keep:
        return head + "…"

    suffix = " … " + " · ".join(fragments_to_keep[:2])
    available = max_len - len(head) - len(suffix)
    if available < 0:
        # Урезаем suffix
        suffix = suffix[:max(10, max_len - len(head))]
    return (head + suffix).strip()


def _priority_sort_key(v: RequirementVerdict) -> tuple:
    return (
        CATEGORY_PRIORITY.get(v.category, 9),
        -v.confidence,
        v.section or "",
        v.requirement_id,
    )


def _assessment_ref_key(assessment: dict | object) -> str:
    source_urls = getattr(assessment, "source_urls", []) if not isinstance(assessment, dict) else assessment.get("source_urls", [])
    source_titles = getattr(assessment, "source_titles", []) if not isinstance(assessment, dict) else assessment.get("source_titles", [])
    if source_urls:
        return str(source_urls[0])
    if source_titles:
        return str(source_titles[0])
    return ""


def _collect_reference_map(report: AnalysisReport) -> dict[str, int]:
    refs: dict[str, int] = {}
    for verdict in report.verdicts:
        for assessment in verdict.platform_assessments:
            key = _assessment_ref_key(assessment)
            if key and key not in refs:
                refs[key] = len(refs) + 1
        for url in verdict.source_urls:
            if url and url not in refs:
                refs[url] = len(refs) + 1
    return refs


def _assessment_ref_label(assessment: dict | object, refs: dict[str, int]) -> str:
    key = _assessment_ref_key(assessment)
    if not key or key not in refs:
        return ""
    return f"[{refs[key]}]"


def _assessment_symbol(assessment: dict | object, refs: dict[str, int]) -> str:
    verdict = getattr(assessment, "verdict", "") if not isinstance(assessment, dict) else assessment.get("verdict", "")
    symbol = PLATFORM_VERDICT_SYMBOLS.get(verdict, "?")
    ref = _assessment_ref_label(assessment, refs)
    return f"{symbol} {ref}".strip()


def _canonical_platform_name(value: str) -> str:
    text = (value or "").strip()
    lowered = text.lower()
    if "гособлак" in lowered or "гос облак" in lowered or "goscloud" in lowered:
        return "ГосОблако"
    if "vmware" in lowered or "vcloud" in lowered or "облако vmware" in lowered:
        return "Облако VMware"
    if "advanced" in lowered:
        return "Advanced"
    if "evolution" in lowered:
        return "Evolution"
    return text


def _is_matrix_platform_name(value: str) -> bool:
    text = _canonical_platform_name(value)
    lowered = text.lower()
    if not text:
        return False
    if lowered.startswith("cloud.ru источник"):
        return False
    if "документация не найдена" in lowered or "платформа не определена" in lowered:
        return False
    return True


_MATRIX_CANONICAL_PLATFORMS = ("ГосОблако", "Облако VMware", "Advanced", "Evolution")


def _is_canonical_cloud_platform(name: str) -> bool:
    """True, если name — каноническая платформа Cloud.ru."""
    return _canonical_platform_name(name) in _MATRIX_CANONICAL_PLATFORMS


def _is_matrix_platform(assessment: object) -> bool:
    raw_name = getattr(assessment, "platform_name", "") or ""
    # external_service отбрасываем ТОЛЬКО если platform_name НЕ каноническая
    # платформа Cloud.ru. Если LLM ошибочно поставил external_service для
    # ГосОблако/VMware/Advanced/Evolution — оставляем его в матрице (имя
    # платформы — ground truth).
    if getattr(assessment, "source_type", "") == "external_service":
        if not _is_canonical_cloud_platform(raw_name):
            return False
    return _is_matrix_platform_name(raw_name)


def _best_platform_assessment(items: list) -> object | None:
    if not items:
        return None
    return sorted(
        items,
        key=lambda item: (
            PLATFORM_VERDICT_RANK.get(getattr(item, "verdict", ""), -1),
            getattr(item, "confidence", 0.0),
        ),
        reverse=True,
    )[0]


def _platform_names(report: AnalysisReport) -> list[str]:
    names = []
    for verdict in report.verdicts:
        for assessment in verdict.platform_assessments:
            if not _is_matrix_platform(assessment):
                continue
            platform_name = _canonical_platform_name(assessment.platform_name)
            if platform_name not in names:
                names.append(platform_name)

    preferred = [name for name in PREFERRED_PLATFORM_ORDER if name in names]
    other = sorted([name for name in names if name not in PREFERRED_PLATFORM_ORDER], key=str.casefold)
    return preferred + other


def _platform_matrix_rows(report: AnalysisReport, refs: dict[str, int]) -> list[dict]:
    platform_names = _platform_names(report)
    rows = []
    for verdict in report.verdicts:
        row = {
            "Пункт ТЗ": verdict.section or f"#{verdict.requirement_id}",
            "Требование": _req_text(verdict, 140),
        }
        by_platform: dict[str, list] = {platform_name: [] for platform_name in platform_names}
        for item in verdict.platform_assessments:
            raw_name = item.platform_name or ""
            platform_name = _canonical_platform_name(raw_name)
            # external_service фильтруется ТОЛЬКО если platform_name НЕ
            # каноническая платформа Cloud.ru (см. _is_matrix_platform).
            if item.source_type == "external_service" and not _is_canonical_cloud_platform(raw_name):
                continue
            if platform_name in by_platform:
                by_platform[platform_name].append(item)
        for platform_name in platform_names:
            assessment = _best_platform_assessment(by_platform.get(platform_name, []))
            # Пустая ячейка = «нет оценки», честнее показать «?» (уточнить),
            # чем «-» (несоответствие). _fill_missing_canonical_platforms на
            # бэкенде должен теперь гарантировать, что 4 канонических
            # ассессмента есть всегда — но fallback оставляем как страховку.
            row[platform_name] = _assessment_symbol(assessment, refs) if assessment else "?"
        rows.append(row)
    return rows


def _platform_totals(report: AnalysisReport) -> dict[str, tuple[int, int, int, int]]:
    totals = {}
    for platform_name in _platform_names(report):
        match_count = partial_count = mismatch_count = clarification_count = 0
        for verdict in report.verdicts:
            # external_service игнорируется только для НЕ-канонических.
            items = [
                assessment for assessment in verdict.platform_assessments
                if (
                    (assessment.source_type != "external_service"
                     or _is_canonical_cloud_platform(assessment.platform_name))
                    and _canonical_platform_name(assessment.platform_name) == platform_name
                )
            ]
            assessment = _best_platform_assessment(items)
            if assessment is None:
                continue
            if assessment.verdict == "match":
                match_count += 1
            elif assessment.verdict == "partial":
                partial_count += 1
            elif assessment.verdict == "mismatch":
                mismatch_count += 1
            else:
                clarification_count += 1
        totals[platform_name] = (match_count, partial_count, mismatch_count, clarification_count)
    return totals


def _reference_title_from_key(key: str) -> str:
    if key.startswith("http"):
        return _url_short_name(key)
    return key


def _decision_summary(report: AnalysisReport) -> str:
    coverage_warnings = _extraction_warning_lines(report)
    if coverage_warnings:
        return "Есть риск неполного анализа ТЗ: " + coverage_warnings[0]
    if report.mismatch_count > 0:
        return "Обнаружены несоответствия. Начните проверку отчёта с блокеров ниже."
    if report.clarification_count > 0:
        return "Критичных блокеров не найдено, но есть пункты, требующие ручного уточнения."
    if report.partial_count > 0:
        return "Явных блокеров не найдено, но есть частичные соответствия, требующие доработки."
    return "Явных блокеров не найдено. Достаточно выборочно перепроверить подтверждённые ключевые требования."


def _extraction_files(report: AnalysisReport) -> list[dict]:
    summary = report.extraction_summary if isinstance(report.extraction_summary, dict) else {}
    files = summary.get("files")
    if isinstance(files, list):
        return [item for item in files if isinstance(item, dict)]
    return [summary] if summary else []


def _extraction_warning_lines(report: AnalysisReport) -> list[str]:
    warnings = []
    for item in _extraction_files(report):
        filename = item.get("filename") or report.document_name
        if item.get("cap_applied"):
            warnings.append(
                f"{filename}: найдено {item.get('requirements_detected_before_cap', '?')} требований, "
                f"в анализ передано {item.get('requirements_returned', '?')}; "
                f"{item.get('requirements_omitted_by_cap', '?')} отброшено лимитом."
            )
        missing = item.get("missing_key_signals_after_extraction") or []
        if missing:
            warnings.append(
                f"{filename}: ключевые сигналы не попали в извлеченные требования: {', '.join(missing[:8])}."
            )
    return warnings


def _extraction_coverage_rows(report: AnalysisReport) -> list[dict]:
    rows = []
    for item in _extraction_files(report):
        filename = item.get("filename") or report.document_name
        rows.append(
            {
                "Файл": filename,
                "Показатель": "Требований найдено парсером",
                "Значение": item.get("requirements_detected_before_cap", ""),
            }
        )
        rows.append(
            {
                "Файл": filename,
                "Показатель": "Передано в анализ",
                "Значение": item.get("requirements_returned", ""),
            }
        )
        rows.append(
            {
                "Файл": filename,
                "Показатель": "Лимит применен",
                "Значение": "да" if item.get("cap_applied") else "нет",
            }
        )
        category_counts = item.get("category_counts_returned") or {}
        if isinstance(category_counts, dict) and category_counts:
            rows.append(
                {
                    "Файл": filename,
                    "Показатель": "Категории в анализе",
                    "Значение": ", ".join(f"{CATEGORY_LABELS.get(k, k)}: {v}" for k, v in category_counts.items()),
                }
            )
        missing = item.get("missing_key_signals_after_extraction") or []
        rows.append(
            {
                "Файл": filename,
                "Показатель": "Ключевые сигналы вне анализа",
                "Значение": ", ".join(missing) if missing else "нет",
            }
        )
    return rows


def _key_signal_rows(report: AnalysisReport) -> list[dict]:
    rows = []
    for item in _extraction_files(report):
        filename = item.get("filename") or report.document_name
        for signal in item.get("key_signal_coverage", []) or []:
            if not isinstance(signal, dict) or not signal.get("present_in_document"):
                continue
            rows.append(
                {
                    "Файл": filename,
                    "Сигнал": signal.get("label", ""),
                    "Критичный": "да" if signal.get("critical") else "нет",
                    "В извлечении": "да" if signal.get("present_in_extracted") else "нет",
                }
            )
    return rows


def _top_key_matches(report: AnalysisReport, limit: int = KEY_MATCHES_LIMIT) -> list[RequirementVerdict]:
    matches = [v for v in report.verdicts if v.verdict == "match"]
    matches.sort(key=_priority_sort_key)
    return matches[:limit]


# ---------------------------------------------------------------------------
# Патч 14 (ZK10). Раздел «Качество анализа»: метрики достоверности отчёта.
# ---------------------------------------------------------------------------

_QUALITY_NORMALIZE_RE = re.compile(r"[^A-Za-zА-Яа-яЁё0-9]+")


def _quality_normalize(text: str) -> str:
    if not text:
        return ""
    return _QUALITY_NORMALIZE_RE.sub(" ", text.lower()).strip()


def _quality_metrics(report: AnalysisReport) -> dict:
    """Собирает технические метрики качества анализа для отчёта (patch 14).

    Сочетает данные post-process (`analysis_quality` в `extraction_summary`)
    с агрегатами, вычисленными прямо по verdict'ам — так старые отчёты,
    в которых post-process ещё не запускался, тоже показывают часть
    информации.
    """
    metrics: dict = {}
    in_scope = [
        v for v in report.verdicts
        if v.verdict != "out_of_scope"
        and (v.category or "").lower() != "procedural"
    ]
    denom = len(in_scope) or 1

    # Дубликаты reasoning'а на 200 нормализованных символах префикса.
    from collections import Counter

    keys = Counter()
    for v in in_scope:
        if not v.reasoning:
            continue
        key = _quality_normalize(v.reasoning[:200])
        if key and len(key) >= 80:
            keys[key] += 1
    duplicate_pairs = sum(c - 1 for c in keys.values() if c >= 2)
    metrics["share_identical_reasoning"] = round(duplicate_pairs / denom, 3)

    metrics["share_confidence_low"] = round(
        sum(1 for v in in_scope if (v.confidence or 0) <= 0.5) / denom, 3
    )
    metrics["share_short_requirements"] = round(
        sum(1 for v in in_scope if len(v.requirement_text or "") <= 25) / denom, 3
    )
    metrics["external_service_match_count"] = sum(
        1 for v in in_scope if v.requires_external_service and v.verdict == "match"
    )

    # Post-process данные (если запускался).
    quality_data = (
        report.extraction_summary.get("analysis_quality")
        if isinstance(report.extraction_summary, dict)
        else None
    ) or {}
    for k, v in quality_data.items():
        metrics.setdefault(k, v)

    # Потерянные / ложно-потерянные сигналы из extraction_summary.files.
    lost: list[dict] = []
    false_pos: list[dict] = []
    files = (
        report.extraction_summary.get("files")
        if isinstance(report.extraction_summary, dict)
        else None
    ) or []
    for item in files:
        if not isinstance(item, dict):
            continue
        filename = item.get("filename") or report.document_name
        for s in item.get("missing_key_signals_after_extraction") or []:
            lost.append({"file": filename, "signal": s})
        for s in item.get("key_signals_false_positives_suspected") or []:
            if isinstance(s, dict):
                false_pos.append({"file": filename, **s})
            else:
                false_pos.append({"file": filename, "label": s})
    metrics["lost_key_signals"] = lost
    metrics["key_signals_false_positives"] = false_pos

    # Patch 15 (strict mode): сборка alert'ов.
    alerts: list[str] = []
    total = report.total or 0
    if total:
        mismatch_ratio = report.mismatch_count / total
        if mismatch_ratio < 0.01:
            alerts.append(
                f"Mismatch floor: mismatch={report.mismatch_count} из {total} "
                f"({mismatch_ratio:.1%}) — анализатор может систематически "
                f"избегать жёстких отказов. Проверьте, нет ли в выборке заведомо "
                f"невыполнимых требований (Tier IV, AWS-specific, КИИ-1 на не-ГосОблаке)."
            )
    if metrics["share_identical_reasoning"] > 0.10:
        alerts.append(
            f"Дубликатов reasoning'а {metrics['share_identical_reasoning']:.1%} — "
            f"возможен template hallucination."
        )
    if metrics["external_service_match_count"] > 0:
        alerts.append(
            f"{metrics['external_service_match_count']} verdict'ов одновременно "
            f"имеют requires_external_service=True и verdict=match — это "
            f"противоречие, patch 9 должен был его убрать."
        )
    if metrics["share_short_requirements"] > 0.10:
        alerts.append(
            f"Доля коротких requirement'ов (≤25 симв.) {metrics['share_short_requirements']:.1%} "
            f"— возможен сбой парсера или утечка table_row-фрагментов."
        )
    metrics["alerts"] = alerts
    return metrics


def _render_quality_section(report: AnalysisReport) -> list[str]:
    m = _quality_metrics(report)
    lines: list[str] = []
    lines.append("## Качество анализа\n")
    lines.append(
        "Технические индикаторы достоверности этого отчёта — выявленные дубликаты "
        "reasoning'а, переиспользованные URL, низкая уверенность, потерянные "
        "ключевые сигналы. Полезно для пресейла: если метрики плохие, отчёт стоит "
        "перепроверить вручную перед отправкой клиенту."
    )
    lines.append("")
    lines.append("| Метрика | Значение |")
    lines.append("|---|---:|")
    if m.get("dedup_reasoning_downgrades") is not None:
        lines.append(f"| Дубликатов reasoning понижено (post-process) | {m['dedup_reasoning_downgrades']} |")
    if m.get("url_overuse_downgrades") is not None:
        lines.append(f"| URL-overuse понижено (post-process) | {m['url_overuse_downgrades']} |")
    if m.get("confidence_snaps") is not None:
        lines.append(f"| Confidence снапнуто (post-process) | {m['confidence_snaps']} |")
    lines.append(f"| Доля одинаковых reasoning'ов | {m['share_identical_reasoning']:.1%} |")
    lines.append(f"| Доля confidence ≤ 0.5 | {m['share_confidence_low']:.1%} |")
    lines.append(f"| Доля requirement'ов ≤ 25 симв. | {m['share_short_requirements']:.1%} |")
    lines.append(f"| external_service+match (должно быть 0) | {m['external_service_match_count']} |")
    lines.append("")

    if m.get("url_overused"):
        lines.append("### URL, переиспользованные больше threshold\n")
        for item in m["url_overused"][:5]:
            url = item.get("url", "")
            uses = item.get("uses", 0)
            lines.append(f"- {url} — {uses} verdict'ов")
        lines.append("")

    if m.get("lost_key_signals"):
        lines.append("### Потерянные ключевые сигналы (direct в документе, нет в извлечении)\n")
        for item in m["lost_key_signals"][:20]:
            lines.append(f"- {item.get('file', '')}: {item.get('signal', '')}")
        lines.append("")

    if m.get("key_signals_false_positives"):
        lines.append("### Возможные ложные потери сигналов (synonym-match без сильного needle)\n")
        for item in m["key_signals_false_positives"][:10]:
            label = item.get("label", "")
            mentions = item.get("total_mentions", "")
            file = item.get("file", "")
            lines.append(
                f"- {file}: {label} (упоминаний в документе: {mentions})"
            )
        lines.append("")

    alerts = m.get("alerts") or []
    if alerts:
        lines.append("### Предупреждения\n")
        for alert in alerts:
            lines.append(f"- ⚠️ {alert}")
        lines.append("")
    return lines


def _md_cell(value) -> str:
    return str(value or "").replace("\n", " ").replace("|", "\\|")


def _suspicious_items(report: AnalysisReport) -> list[dict]:
    return report.suspicious_items


def _trace_rows(report: AnalysisReport) -> list[dict]:
    rows = []
    for verdict in report.verdicts:
        trace = verdict.trace or {}
        profile = trace.get("profile", {}) if isinstance(trace.get("profile"), dict) else {}
        selected_sources = trace.get("selected_sources", []) if isinstance(trace.get("selected_sources"), list) else []
        if not selected_sources:
            rows.append(
                {
                    "Пункт ТЗ": verdict.section or f"#{verdict.requirement_id}",
                    "Профиль": profile.get("cluster", ""),
                    "Платформа": profile.get("platform_hint", ""),
                    "Источник": "",
                    "Score": "",
                    "Причины выбора": trace.get("rag_error", "нет выбранных источников"),
                    "Вердикт": VERDICT_LABELS.get(verdict.verdict, verdict.verdict),
                    "Evidence status": verdict.evidence_status,
                }
            )
            continue
        for source in selected_sources[:3]:
            rows.append(
                {
                    "Пункт ТЗ": verdict.section or f"#{verdict.requirement_id}",
                    "Профиль": profile.get("cluster", ""),
                    "Платформа": source.get("platform") or profile.get("platform_hint", ""),
                    "Источник": source.get("title", ""),
                    "Score": source.get("score", ""),
                    "Причины выбора": "; ".join(source.get("reasons", []) or []),
                    "Вердикт": VERDICT_LABELS.get(verdict.verdict, verdict.verdict),
                    "Evidence status": verdict.evidence_status,
                }
            )
    return rows


def _format_problem_entry(lines: list[str], verdict: RequirementVerdict, reason_label: str) -> None:
    cat = CATEGORY_LABELS.get(verdict.category, verdict.category)
    lines.append(f"### {_section_label(verdict)} ({cat})")
    lines.append(f"**Пункт / раздел в ТЗ:** {verdict.section or f'#{verdict.requirement_id}'}")
    lines.append(f"**Требование:** {_req_text(verdict)}")
    lines.append(f"**{reason_label}:** {verdict.reasoning or 'Требуется ручная проверка'}")
    if verdict.recommendation:
        lines.append(f"**Рекомендация:** {verdict.recommendation}")
    if verdict.source_urls:
        links = ", ".join(f"[{_url_short_name(u)}]({u})" for u in verdict.source_urls[:5])
        lines.append(f"**Документация:** {links}")
    if verdict.platform_assessments:
        lines.append("**Оценка по платформам/услугам:**")
        for assessment in verdict.platform_assessments:
            label = VERDICT_LABELS.get(assessment.verdict, assessment.verdict)
            source_type = "внешняя услуга" if assessment.source_type == "external_service" else "платформа"
            lines.append(f"- {assessment.platform_name} ({source_type}): {label}. {assessment.reasoning}")
    if verdict.requires_external_service:
        lines.append(f"**Нужна проработка подрядчиков:** {verdict.external_service_notes or 'Да'}")
    lines.append("")


def generate_markdown(report: AnalysisReport) -> str:
    """Generate a full Markdown report."""
    lines = []
    now = datetime.now().strftime("%d.%m.%Y %H:%M")

    lines.append(f"# Отчёт по анализу ТЗ: {report.document_name}")
    lines.append(f"*Дата анализа: {now}*\n")

    # Prominent compliance percentage at the very top — главная цифра для
    # пресейла теперь «Покрытие на рекомендуемой платформе». Это процент по
    # platform_assessments выбранной канонической платформы Cloud.ru, а не
    # средневзвешенный по всем требованиям. Старая метрика
    # compliance_percentage остаётся ниже — для обратной совместимости и как
    # «средний best-case» по портфелю.
    rec_platform = report.recommended_platform
    rec_pct = report.recommended_platform_compliance
    pct = report.compliance_percentage
    if rec_platform:
        lines.append(
            f"## Покрытие на рекомендуемой платформе ({rec_platform}): {rec_pct}%"
        )
        lines.append(
            f"*Это доля требований ТЗ, которые закрываются на платформе "
            f"{rec_platform} — именно её Cloud.ru предлагает заказчику.*\n"
        )
        lines.append(f"### Общий процент соответствия портфеля: {pct}%")
        lines.append(
            f"*Best-case по любой платформе Cloud.ru: {report.score} из "
            f"{report.max_score} баллов.*\n"
        )
    else:
        lines.append(f"## Общий процент соответствия: {pct}%")
        lines.append(f"**{report.score} из {report.max_score} баллов**\n")
    lines.append(f"**Быстрый вывод:** {_decision_summary(report)}\n")

    # Methodology
    lines.append("### Методика оценки\n")
    lines.append("| Вердикт | Баллы |")
    lines.append("|---|---|")
    lines.append("| Полное соответствие | 2 |")
    lines.append("| Частичное соответствие | 1 |")
    lines.append("| Несоответствие / Требует уточнения | 0 |")
    lines.append("")
    lines.append(f"Максимальный балл = {report.total} пунктов × 2 = {report.max_score}. "
                 f"Набрано {report.score} баллов. "
                 f"Итоговый процент = {report.score} / {report.max_score} × 100 = **{pct}%**")
    lines.append("")

    # Summary table
    lines.append("## Сводка\n")
    lines.append(f"| Показатель | Значение |")
    lines.append(f"|---|---|")
    lines.append(f"| Всего требований | {report.total} |")
    lines.append(f"| {VERDICT_ICONS['match']} Соответствует | {report.match_count} |")
    lines.append(f"| {VERDICT_ICONS['partial']} Частично | {report.partial_count} |")
    lines.append(f"| {VERDICT_ICONS['mismatch']} Не соответствует | {report.mismatch_count} |")
    lines.append(f"| {VERDICT_ICONS['needs_clarification']} Требует уточнения | {report.clarification_count} |")
    lines.append(f"| **Баллы** | **{report.score} / {report.max_score}** |")
    lines.append(f"| **Общее соответствие (best-case)** | **{pct}%** |")
    if rec_platform:
        lines.append(
            f"| **Покрытие {rec_platform}** | **{rec_pct}%** |"
        )
    lines.append("")

    coverage_rows = _extraction_coverage_rows(report)
    signal_rows = _key_signal_rows(report)
    if coverage_rows:
        lines.append("## Покрытие извлечения\n")
        warning_lines = _extraction_warning_lines(report)
        if warning_lines:
            lines.append("**Внимание:**")
            for warning in warning_lines:
                lines.append(f"- {warning}")
            lines.append("")
        lines.append("| Файл | Показатель | Значение |")
        lines.append("|---|---|---|")
        for row in coverage_rows:
            lines.append(
                f"| {_md_cell(row.get('Файл'))} | {_md_cell(row.get('Показатель'))} | {_md_cell(row.get('Значение'))} |"
            )
        lines.append("")
        if signal_rows:
            lines.append("| Файл | Ключевой сигнал | Критичный | Попал в извлечение |")
            lines.append("|---|---|---|---|")
            for row in signal_rows:
                lines.append(
                    f"| {_md_cell(row.get('Файл'))} | {_md_cell(row.get('Сигнал'))} | "
                    f"{_md_cell(row.get('Критичный'))} | {_md_cell(row.get('В извлечении'))} |"
                )
            lines.append("")

    if report.summary:
        lines.append("### Резюме\n")
        lines.append(report.summary)
        lines.append("")

    refs = _collect_reference_map(report)
    matrix_rows = _platform_matrix_rows(report, refs)
    platform_names = _platform_names(report)
    if matrix_rows and platform_names:
        lines.append("## Матрица соответствия по платформам\n")
        header = ["Пункт ТЗ", "Требование"] + platform_names
        lines.append("| " + " | ".join(header) + " |")
        lines.append("| " + " | ".join(["---"] * len(header)) + " |")
        for row in matrix_rows:
            lines.append("| " + " | ".join(str(row.get(column, "")).replace("\n", " ") for column in header) + " |")
        totals = _platform_totals(report)
        lines.append("")
        lines.append("### Итого по платформам\n")
        lines.append("| Платформа / услуга | + | ± | - | ? |")
        lines.append("|---|---:|---:|---:|---:|")
        for platform_name, (match_count, partial_count, mismatch_count, clarification_count) in totals.items():
            lines.append(f"| {platform_name} | {match_count} | {partial_count} | {mismatch_count} | {clarification_count} |")
        lines.append("")

    external_items = [v for v in report.verdicts if v.requires_external_service]
    if external_items:
        lines.append("## Требования для проработки внешних услуг / подрядчиков\n")
        for v in external_items:
            lines.append(f"- **{v.section or f'#{v.requirement_id}'}**: {_req_text(v, 180)}")
            if v.external_service_notes:
                lines.append(f"  - {v.external_service_notes}")
        lines.append("")

    suspicious_items = _suspicious_items(report)
    if suspicious_items:
        lines.append("## Сомнительные места\n")
        lines.append("| Пункт ТЗ | Вердикт | Уверенность | Причины | Требование |")
        lines.append("|---|---|---:|---|---|")
        for item in suspicious_items:
            lines.append(
                "| "
                + " | ".join(
                    [
                        _md_cell(item.get("section") or f"#{item.get('requirement_id')}"),
                        _md_cell(VERDICT_LABELS.get(item.get("verdict", ""), item.get("verdict", ""))),
                        _md_cell(f"{float(item.get('confidence') or 0):.0%}"),
                        _md_cell("; ".join(item.get("reasons", []) or [])),
                        _md_cell(str(item.get("requirement_text", ""))[:180]),
                    ]
                )
                + " |"
            )
        lines.append("")

    trace_rows = _trace_rows(report)
    if trace_rows:
        lines.append("## Трассировка RAG\n")
        lines.append("| Пункт ТЗ | Профиль | Платформа | Источник | Score | Причины выбора | Evidence |")
        lines.append("|---|---|---|---|---:|---|---|")
        for row in trace_rows:
            lines.append(
                "| "
                + " | ".join(
                    [
                        _md_cell(row.get("Пункт ТЗ")),
                        _md_cell(row.get("Профиль")),
                        _md_cell(row.get("Платформа")),
                        _md_cell(row.get("Источник")),
                        _md_cell(row.get("Score")),
                        _md_cell(row.get("Причины выбора")),
                        _md_cell(row.get("Evidence status")),
                    ]
                )
                + " |"
            )
        lines.append("")

    mismatches = sorted([v for v in report.verdicts if v.verdict == "mismatch"], key=_priority_sort_key)
    clarifications = sorted([v for v in report.verdicts if v.verdict == "needs_clarification"], key=_priority_sort_key)
    partials = sorted([v for v in report.verdicts if v.verdict == "partial"], key=_priority_sort_key)
    key_matches = _top_key_matches(report)

    # Patch 14 (ZK10): секция «Качество анализа» — индикаторы достоверности.
    lines.extend(_render_quality_section(report))

    # Priority checks first
    lines.append("## Что проверить в первую очередь\n")
    lines.append(f"- Несоответствия: **{report.mismatch_count}**")
    lines.append(f"- Требуют уточнения: **{report.clarification_count}**")
    lines.append(f"- Частичные соответствия: **{report.partial_count}**")
    lines.append("")

    if mismatches:
        lines.append("## Несоответствия\n")
        for v in mismatches:
            _format_problem_entry(lines, v, "Причина")

    if clarifications:
        lines.append("## Требуют уточнения\n")
        for v in clarifications:
            _format_problem_entry(lines, v, "Комментарий")

    if partials:
        lines.append("## Частичное соответствие\n")
        for v in partials:
            _format_problem_entry(lines, v, "Обоснование")

    if key_matches:
        lines.append("## Подтверждённые важные соответствия\n")
        lines.append(f"Показаны только наиболее значимые подтверждённые пункты: {len(key_matches)} из {report.match_count}.")
        lines.append("")
        for v in key_matches:
            cat = CATEGORY_LABELS.get(v.category, v.category)
            lines.append(f"### {_section_label(v)} ({cat})")
            lines.append(f"**Пункт / раздел в ТЗ:** {v.section or f'#{v.requirement_id}'}")
            lines.append(f"**Требование:** {_req_text(v)}")
            if v.reasoning:
                lines.append(f"**Почему соответствует:** {v.reasoning}")
            if v.source_urls:
                links = ", ".join(f"[{_url_short_name(u)}]({u})" for u in v.source_urls[:5])
                lines.append(f"**Документация:** {links}")
            lines.append("")

    # Процедурные пункты закупки — вне технической оценки, но не пропали.
    procedural = sorted(
        [v for v in report.verdicts if (v.category or "").lower() == "procedural"
         or v.verdict == "out_of_scope"],
        key=lambda v: v.requirement_id,
    )
    if procedural:
        lines.append("## Процедурные пункты закупки (вне технической оценки)\n")
        lines.append(
            f"Извлечено {len(procedural)} пунктов, относящихся к коммерческо-правовой "
            "обвязке тендера (ОКПД, начальная максимальная цена, обеспечение заявки, "
            "антикоррупция, идентификация участников закупки, реквизиты сторон). "
            "Они **не оценивают технические возможности Cloud.ru** и исключены из "
            "знаменателя процента соответствия. Передайте их в коммерческую/правовую "
            "команду Cloud.ru при подготовке КП."
        )
        lines.append("")
        lines.append("| Пункт ТЗ | Требование |")
        lines.append("|---|---|")
        for v in procedural:
            lines.append(f"| {_md_cell(v.section or f'#{v.requirement_id}')} | {_md_cell(_req_text(v, max_len=300))} |")
        lines.append("")

    if refs:
        lines.append("## Сноски RAG\n")
        for key, index in sorted(refs.items(), key=lambda item: item[1]):
            if key.startswith("http"):
                lines.append(f"[{index}] [{_reference_title_from_key(key)}]({key})")
            else:
                lines.append(f"[{index}] {key}")
        lines.append("")

    # Full detail by category (без procedural — у них отдельный раздел выше)
    lines.append("## Детализация по всем требованиям\n")
    categories_order = ["technical", "sla", "security", "legal", "commercial", "other"]
    for cat_key in categories_order:
        cat_verdicts = [v for v in report.verdicts if v.category == cat_key]
        if not cat_verdicts:
            continue
        cat_label = CATEGORY_LABELS.get(cat_key, cat_key)
        lines.append(f"### {cat_label}\n")
        for v in cat_verdicts:
            icon = VERDICT_ICONS.get(v.verdict, "")
            label = VERDICT_LABELS.get(v.verdict, v.verdict)
            lines.append(f"**{_section_label(v)}** {icon} {label} (уверенность: {v.confidence:.0%})")
            lines.append(f"> {_req_text(v)}")
            if v.reasoning:
                lines.append(f"\n*Обоснование:* {v.reasoning}")
            if v.evidence:
                lines.append(f"\n*Источник:* {v.evidence}")
            if v.source_urls:
                links = ", ".join(f"[{_url_short_name(u)}]({u})" for u in v.source_urls[:5])
                lines.append(f"\n*Документация:* {links}")
            if v.recommendation:
                lines.append(f"\n*Рекомендация:* {v.recommendation}")
            lines.append("")

    return "\n".join(lines)


def _url_short_name(url: str) -> str:
    """Extract a short label from a cloud.ru/docs URL."""
    # https://cloud.ru/docs/s3e/ug/topics/overview.html -> s3e / overview
    path = url.replace("https://cloud.ru/docs/", "").rstrip("/")
    parts = [p for p in path.split("/") if p not in ("ug", "topics", "index", "index.html")]
    if len(parts) > 2:
        return f"{parts[0]} / {parts[-1]}"
    return " / ".join(parts) if parts else url


def save_markdown(report: AnalysisReport, output_dir: Path | None = None) -> Path:
    """Save report as Markdown file."""
    output_dir = output_dir or cfg.REPORTS_DIR
    output_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in report.document_name)
    filename = f"report_{safe_name}_{timestamp}.md"
    path = output_dir / filename
    md = generate_markdown(report)
    path.write_text(md, encoding="utf-8")
    logger.info("Saved Markdown report: %s", path)
    return path


def save_docx(report: AnalysisReport, output_dir: Path | None = None) -> Path:
    """Save report as DOCX file."""
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    output_dir = output_dir or cfg.REPORTS_DIR
    output_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in report.document_name)
    filename = f"report_{safe_name}_{timestamp}.docx"
    path = output_dir / filename

    doc = Document()

    # Title
    doc.add_heading(f"Отчёт по анализу ТЗ: {report.document_name}", level=0)
    doc.add_paragraph(f"Дата анализа: {datetime.now().strftime('%d.%m.%Y %H:%M')}")

    # Prominent compliance percentage
    pct = report.compliance_percentage
    p = doc.add_paragraph()
    run = p.add_run(f"Общий процент соответствия: {pct}%")
    run.bold = True
    run.font.size = Pt(18)
    p = doc.add_paragraph()
    run = p.add_run(f"{report.score} из {report.max_score} баллов")
    run.bold = True
    run.font.size = Pt(14)

    # Methodology
    doc.add_heading("Методика оценки", level=2)
    meth_table = doc.add_table(rows=4, cols=2)
    meth_table.style = "Table Grid"
    meth_data = [
        ("Вердикт", "Баллы"),
        ("Полное соответствие", "2"),
        ("Частичное соответствие", "1"),
        ("Несоответствие / Требует уточнения", "0"),
    ]
    for i, (k, val) in enumerate(meth_data):
        meth_table.rows[i].cells[0].text = k
        meth_table.rows[i].cells[1].text = val
    doc.add_paragraph(
        f"Максимальный балл = {report.total} пунктов × 2 = {report.max_score}. "
        f"Набрано {report.score} баллов. "
        f"Итоговый процент = {report.score} / {report.max_score} × 100 = {pct}%"
    )

    # Summary table
    doc.add_heading("Сводка", level=1)
    table = doc.add_table(rows=8, cols=2)
    table.style = "Table Grid"
    summary_data = [
        ("Показатель", "Значение"),
        ("Всего требований", str(report.total)),
        ("Соответствует", str(report.match_count)),
        ("Частично", str(report.partial_count)),
        ("Не соответствует", str(report.mismatch_count)),
        ("Требует уточнения", str(report.clarification_count)),
        ("Баллы", f"{report.score} / {report.max_score}"),
        ("Общее соответствие", f"{pct}%"),
    ]
    for i, (k, v) in enumerate(summary_data):
        table.rows[i].cells[0].text = k
        table.rows[i].cells[1].text = v

    coverage_rows = _extraction_coverage_rows(report)
    signal_rows = _key_signal_rows(report)
    if coverage_rows:
        doc.add_heading("Покрытие извлечения", level=1)
        for warning in _extraction_warning_lines(report):
            doc.add_paragraph(f"Внимание: {warning}")
        coverage_table = doc.add_table(rows=1, cols=3)
        coverage_table.style = "Table Grid"
        for j, h in enumerate(["Файл", "Показатель", "Значение"]):
            coverage_table.rows[0].cells[j].text = h
        for item in coverage_rows:
            row = coverage_table.add_row()
            row.cells[0].text = str(item.get("Файл", ""))
            row.cells[1].text = str(item.get("Показатель", ""))
            row.cells[2].text = str(item.get("Значение", ""))
        if signal_rows:
            signals_table = doc.add_table(rows=1, cols=4)
            signals_table.style = "Table Grid"
            for j, h in enumerate(["Файл", "Ключевой сигнал", "Критичный", "Попал в извлечение"]):
                signals_table.rows[0].cells[j].text = h
            for item in signal_rows:
                row = signals_table.add_row()
                row.cells[0].text = str(item.get("Файл", ""))
                row.cells[1].text = str(item.get("Сигнал", ""))
                row.cells[2].text = str(item.get("Критичный", ""))
                row.cells[3].text = str(item.get("В извлечении", ""))

    if report.summary:
        doc.add_heading("Резюме", level=2)
        doc.add_paragraph(report.summary)

    refs = _collect_reference_map(report)
    matrix_rows = _platform_matrix_rows(report, refs)
    platform_names = _platform_names(report)
    if matrix_rows and platform_names:
        doc.add_heading("Матрица соответствия по платформам", level=1)
        header = ["Пункт ТЗ", "Требование"] + platform_names
        matrix_table = doc.add_table(rows=1, cols=len(header))
        matrix_table.style = "Table Grid"
        for j, h in enumerate(header):
            matrix_table.rows[0].cells[j].text = h
        for item in matrix_rows:
            row = matrix_table.add_row()
            for j, column in enumerate(header):
                row.cells[j].text = str(item.get(column, ""))

        doc.add_heading("Итого по платформам", level=2)
        totals = _platform_totals(report)
        totals_table = doc.add_table(rows=1, cols=5)
        totals_table.style = "Table Grid"
        for j, h in enumerate(["Платформа / услуга", "+", "±", "-", "?"]):
            totals_table.rows[0].cells[j].text = h
        for platform_name, (match_count, partial_count, mismatch_count, clarification_count) in totals.items():
            row = totals_table.add_row()
            row.cells[0].text = platform_name
            row.cells[1].text = str(match_count)
            row.cells[2].text = str(partial_count)
            row.cells[3].text = str(mismatch_count)
            row.cells[4].text = str(clarification_count)

    external_items = [v for v in report.verdicts if v.requires_external_service]
    if external_items:
        doc.add_heading("Требования для проработки внешних услуг / подрядчиков", level=1)
        for v in external_items:
            doc.add_paragraph(
                f"{v.section or f'#{v.requirement_id}'}: {_req_text(v, 180)} "
                f"{v.external_service_notes or ''}",
                style=None,
            )

    suspicious_items = _suspicious_items(report)
    if suspicious_items:
        doc.add_heading("Сомнительные места", level=1)
        t = doc.add_table(rows=1, cols=5)
        t.style = "Table Grid"
        for j, h in enumerate(["Пункт ТЗ", "Вердикт", "Уверенность", "Причины", "Требование"]):
            t.rows[0].cells[j].text = h
        for item in suspicious_items:
            row = t.add_row()
            row.cells[0].text = str(item.get("section") or f"#{item.get('requirement_id')}")
            row.cells[1].text = VERDICT_LABELS.get(item.get("verdict", ""), item.get("verdict", ""))
            row.cells[2].text = f"{float(item.get('confidence') or 0):.0%}"
            row.cells[3].text = "; ".join(item.get("reasons", []) or [])
            row.cells[4].text = str(item.get("requirement_text", ""))[:180]

    trace_rows = _trace_rows(report)
    if trace_rows:
        doc.add_heading("Трассировка RAG", level=1)
        t = doc.add_table(rows=1, cols=6)
        t.style = "Table Grid"
        for j, h in enumerate(["Пункт ТЗ", "Профиль", "Платформа", "Источник", "Score", "Причины выбора"]):
            t.rows[0].cells[j].text = h
        for item in trace_rows[:120]:
            row = t.add_row()
            row.cells[0].text = str(item.get("Пункт ТЗ", ""))
            row.cells[1].text = str(item.get("Профиль", ""))
            row.cells[2].text = str(item.get("Платформа", ""))
            row.cells[3].text = str(item.get("Источник", ""))[:180]
            row.cells[4].text = str(item.get("Score", ""))
            row.cells[5].text = str(item.get("Причины выбора", ""))[:240]

    doc.add_heading("Что проверить в первую очередь", level=1)
    doc.add_paragraph(_decision_summary(report))

    mismatches = sorted([v for v in report.verdicts if v.verdict == "mismatch"], key=_priority_sort_key)
    clarifications = sorted([v for v in report.verdicts if v.verdict == "needs_clarification"], key=_priority_sort_key)
    partials = sorted([v for v in report.verdicts if v.verdict == "partial"], key=_priority_sort_key)
    key_matches = _top_key_matches(report)

    if mismatches:
        doc.add_heading("Несоответствия", level=1)
        t = doc.add_table(rows=1, cols=6)
        t.style = "Table Grid"
        for j, h in enumerate(["Пункт ТЗ", "Категория", "Требование", "Причина", "Рекомендация", "Ссылки на документацию"]):
            t.rows[0].cells[j].text = h
        for v in mismatches:
            row = t.add_row()
            row.cells[0].text = v.section or f"#{v.requirement_id}"
            row.cells[1].text = CATEGORY_LABELS.get(v.category, v.category)
            row.cells[2].text = _req_text(v, 150)
            row.cells[3].text = v.reasoning
            row.cells[4].text = v.recommendation
            row.cells[5].text = "\n".join(v.source_urls[:5]) if v.source_urls else ""

    if clarifications:
        doc.add_heading("Требуют уточнения", level=1)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for j, h in enumerate(["Пункт ТЗ", "Категория", "Требование", "Комментарий"]):
            t.rows[0].cells[j].text = h
        for v in clarifications:
            row = t.add_row()
            row.cells[0].text = v.section or f"#{v.requirement_id}"
            row.cells[1].text = CATEGORY_LABELS.get(v.category, v.category)
            row.cells[2].text = _req_text(v, 150)
            row.cells[3].text = v.reasoning

    if partials:
        doc.add_heading("Частичное соответствие", level=1)
        t = doc.add_table(rows=1, cols=6)
        t.style = "Table Grid"
        for j, h in enumerate(["Пункт ТЗ", "Категория", "Требование", "Обоснование", "Рекомендация", "Ссылки на документацию"]):
            t.rows[0].cells[j].text = h
        for v in partials:
            row = t.add_row()
            row.cells[0].text = v.section or f"#{v.requirement_id}"
            row.cells[1].text = CATEGORY_LABELS.get(v.category, v.category)
            row.cells[2].text = _req_text(v, 150)
            row.cells[3].text = v.reasoning
            row.cells[4].text = v.recommendation
            row.cells[5].text = "\n".join(v.source_urls[:5]) if v.source_urls else ""

    if key_matches:
        doc.add_heading("Подтверждённые важные соответствия", level=1)
        t = doc.add_table(rows=1, cols=4)
        t.style = "Table Grid"
        for j, h in enumerate(["Пункт ТЗ", "Категория", "Требование", "Документация"]):
            t.rows[0].cells[j].text = h
        for v in key_matches:
            row = t.add_row()
            row.cells[0].text = v.section or f"#{v.requirement_id}"
            row.cells[1].text = CATEGORY_LABELS.get(v.category, v.category)
            row.cells[2].text = _req_text(v, 150)
            row.cells[3].text = "\n".join(v.source_urls[:5]) if v.source_urls else ""

    # Процедурные пункты закупки — вне технической оценки Cloud.ru.
    procedural = sorted(
        [v for v in report.verdicts if (v.category or "").lower() == "procedural"
         or v.verdict == "out_of_scope"],
        key=lambda v: v.requirement_id,
    )
    if procedural:
        doc.add_heading("Процедурные пункты закупки (вне технической оценки)", level=1)
        doc.add_paragraph(
            f"Извлечено {len(procedural)} пунктов, относящихся к коммерческо-правовой "
            "обвязке тендера (ОКПД, начальная максимальная цена, обеспечение заявки, "
            "антикоррупция, идентификация участников закупки, реквизиты сторон). "
            "Они не оценивают технические возможности Cloud.ru и исключены из "
            "знаменателя процента соответствия. Передайте их в коммерческую/правовую "
            "команду Cloud.ru при подготовке КП."
        )
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        for j, h in enumerate(["Пункт ТЗ", "Требование"]):
            t.rows[0].cells[j].text = h
        for v in procedural:
            row = t.add_row()
            row.cells[0].text = v.section or f"#{v.requirement_id}"
            row.cells[1].text = _req_text(v, 300)

    if refs:
        doc.add_heading("Сноски RAG", level=1)
        for key, index in sorted(refs.items(), key=lambda item: item[1]):
            doc.add_paragraph(f"[{index}] {key}")

    doc.save(str(path))
    logger.info("Saved DOCX report: %s", path)
    return path


def save_excel(report: AnalysisReport, output_dir: Path | None = None) -> Path:
    """Save report as Excel file with a detailed table of all requirements."""
    import pandas as pd

    output_dir = output_dir or cfg.REPORTS_DIR
    output_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in report.document_name)
    filename = f"report_{safe_name}_{timestamp}.xlsx"
    path = output_dir / filename

    # Build rows for the main table
    rows = []
    for v in report.verdicts:
        score_val = 2 if v.verdict == "match" else (1 if v.verdict == "partial" else 0)
        platform_summary = []
        for assessment in v.platform_assessments:
            platform_summary.append(
                f"{assessment.platform_name}: {VERDICT_LABELS.get(assessment.verdict, assessment.verdict)}"
            )
        rows.append({
            "Пункт ТЗ": v.section or f"#{v.requirement_id}",
            "Текст требования": _req_text(v, 1000),
            "Категория": CATEGORY_LABELS.get(v.category, v.category),
            "Вердикт": VERDICT_LABELS.get(v.verdict, v.verdict),
            "Баллы": score_val,
            "Уверенность": f"{v.confidence:.0%}",
            "Обоснование": v.reasoning,
            "Источник (цитата)": v.evidence,
            "Рекомендация": v.recommendation,
            "Ссылки на документацию": "\n".join(v.source_urls[:5]) if v.source_urls else "",
            "Оценка по платформам": "\n".join(platform_summary),
            "Нужна проработка подрядчиков": "Да" if v.requires_external_service else "Нет",
            "Внешние услуги / комментарий": v.external_service_notes,
            "Evidence status": v.evidence_status,
            "Evidence contract notes": "\n".join(v.evidence_contract_notes),
        })

    df = pd.DataFrame(rows)
    refs = _collect_reference_map(report)
    df_matrix = pd.DataFrame(_platform_matrix_rows(report, refs))
    detail_rows = []
    for v in report.verdicts:
        for assessment in v.platform_assessments:
            detail_rows.append(
                {
                    "Пункт ТЗ": v.section or f"#{v.requirement_id}",
                    "Требование": _req_text(v, 500),
                    "Платформа / услуга": assessment.platform_name,
                    "Тип источника": assessment.source_type,
                    "Вердикт": VERDICT_LABELS.get(assessment.verdict, assessment.verdict),
                    "Уверенность": f"{assessment.confidence:.0%}",
                    "Обоснование": assessment.reasoning,
                    "Сноски": ", ".join(assessment.evidence_refs),
                    "Источники": "\n".join(assessment.source_urls or assessment.source_titles),
                    "Рекомендация": assessment.recommendation,
                }
            )
    df_platform_detail = pd.DataFrame(detail_rows)
    df_refs = pd.DataFrame(
        [{"Сноска": f"[{index}]", "Источник": key} for key, index in sorted(refs.items(), key=lambda item: item[1])]
    )
    df_suspicious = pd.DataFrame(
        [
            {
                "Пункт ТЗ": item.get("section") or f"#{item.get('requirement_id')}",
                "Требование": str(item.get("requirement_text", ""))[:500],
                "Вердикт": VERDICT_LABELS.get(item.get("verdict", ""), item.get("verdict", "")),
                "Уверенность": f"{float(item.get('confidence') or 0):.0%}",
                "Причины": "\n".join(item.get("reasons", []) or []),
                "Рекомендация": item.get("recommendation", ""),
            }
            for item in _suspicious_items(report)
        ]
    )
    df_trace = pd.DataFrame(_trace_rows(report))
    df_extraction = pd.DataFrame(_extraction_coverage_rows(report))
    df_signals = pd.DataFrame(_key_signal_rows(report))

    pct = report.compliance_percentage

    # Summary data
    summary_rows = [
        {"Показатель": "Всего требований", "Значение": report.total},
        {"Показатель": "Соответствует", "Значение": report.match_count},
        {"Показатель": "Частично", "Значение": report.partial_count},
        {"Показатель": "Не соответствует", "Значение": report.mismatch_count},
        {"Показатель": "Требует уточнения", "Значение": report.clarification_count},
        {"Показатель": "Баллы", "Значение": f"{report.score} / {report.max_score}"},
        {"Показатель": "Общее соответствие", "Значение": f"{pct}%"},
        {"Показатель": "", "Значение": ""},
        {"Показатель": "Методика оценки", "Значение": ""},
        {"Показатель": "Полное соответствие", "Значение": "2 балла"},
        {"Показатель": "Частичное соответствие", "Значение": "1 балл"},
        {"Показатель": "Несоответствие / Требует уточнения", "Значение": "0 баллов"},
    ]
    df_summary = pd.DataFrame(summary_rows)

    with pd.ExcelWriter(str(path), engine="openpyxl") as writer:
        df_summary.to_excel(writer, sheet_name="Сводка", index=False)
        if not df_matrix.empty:
            df_matrix.to_excel(writer, sheet_name="Матрица платформ", index=False)
        if not df_platform_detail.empty:
            df_platform_detail.to_excel(writer, sheet_name="Платформы детально", index=False)
        if not df_refs.empty:
            df_refs.to_excel(writer, sheet_name="Сноски RAG", index=False)
        if not df_suspicious.empty:
            df_suspicious.to_excel(writer, sheet_name="Сомнительные места", index=False)
        if not df_trace.empty:
            df_trace.to_excel(writer, sheet_name="Трассировка RAG", index=False)
        if not df_extraction.empty:
            df_extraction.to_excel(writer, sheet_name="Покрытие извлечения", index=False)
        if not df_signals.empty:
            df_signals.to_excel(writer, sheet_name="Ключевые сигналы", index=False)
        df.to_excel(writer, sheet_name="Все требования", index=False)

        # Separate sheets by verdict
        for verdict_key, label in VERDICT_LABELS.items():
            filtered = df[df["Вердикт"] == label]
            if not filtered.empty:
                sheet_name = label[:31]  # Excel sheet name max 31 chars
                filtered.to_excel(writer, sheet_name=sheet_name, index=False)

        # Auto-adjust column widths
        for sheet_name in writer.sheets:
            ws = writer.sheets[sheet_name]
            for col_cells in ws.columns:
                max_len = 0
                col_letter = col_cells[0].column_letter
                for cell in col_cells:
                    try:
                        cell_len = len(str(cell.value or ""))
                        if cell_len > max_len:
                            max_len = cell_len
                    except Exception:
                        pass
                ws.column_dimensions[col_letter].width = min(max_len + 2, 60)

    logger.info("Saved Excel report: %s", path)
    return path


def save_pdf(report: AnalysisReport, output_dir: Path | None = None) -> Path:
    """Save report as PDF file using fpdf2."""
    from fpdf import FPDF, XPos, YPos

    output_dir = output_dir or cfg.REPORTS_DIR
    output_dir.mkdir(exist_ok=True)
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    safe_name = "".join(c if c.isalnum() or c in "._-" else "_" for c in report.document_name)
    filename = f"report_{safe_name}_{timestamp}.pdf"
    path = output_dir / filename

    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)

    # Try to add a Unicode font with Cyrillic support
    font_added = False
    font_paths = [
        "/Library/Fonts/Arial Unicode.ttf",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf",
        "/usr/share/fonts/truetype/liberation/LiberationSans-Regular.ttf",
    ]
    for fp in font_paths:
        if Path(fp).exists():
            try:
                pdf.add_font("UniFont", "", fp)
                font_added = True
                break
            except Exception:
                continue

    if not font_added:
        import glob
        for pattern in ["/Library/Fonts/*.ttf", "/System/Library/Fonts/*.ttf",
                        "/System/Library/Fonts/Supplemental/*.ttf"]:
            for f in glob.glob(pattern):
                try:
                    pdf.add_font("UniFont", "", f)
                    font_added = True
                    break
                except Exception:
                    continue
            if font_added:
                break

    font_name = "UniFont" if font_added else "Helvetica"
    # UniFont doesn't have a bold variant; we'll simulate bold via font size

    def _safe(text: str) -> str:
        """Clean text for PDF output."""
        return text.replace("\r", "").strip()

    def _cut(text: str, limit: int) -> str:
        text = _safe(str(text or "")).replace("\n", " ")
        return text if len(text) <= limit else text[: limit - 1] + "…"

    pdf.add_page()

    # Title
    # Title
    pdf.set_font(font_name, "", 16)
    pdf.multi_cell(0, 10, _safe(f"Отчет по анализу ТЗ: {report.document_name}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.set_font(font_name, "", 10)
    pdf.cell(0, 8, f"Дата анализа: {datetime.now().strftime('%d.%m.%Y %H:%M')}", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Compliance percentage — prominent
    pct = report.compliance_percentage
    pdf.set_font(font_name, "", 18)
    pdf.cell(0, 12, _safe(f"Общий процент соответствия: {pct}%"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(font_name, "", 14)
    pdf.cell(0, 10, _safe(f"{report.score} из {report.max_score} баллов"), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Methodology
    pdf.set_font(font_name, "", 12)
    pdf.cell(0, 8, _safe("Методика оценки"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(font_name, "", 10)
    for line in [
        "Полное соответствие = 2 балла",
        "Частичное соответствие = 1 балл",
        "Несоответствие / Требует уточнения = 0 баллов",
        f"Максимальный балл = {report.total} x 2 = {report.max_score}",
        f"Набрано {report.score} баллов. Итого: {pct}%",
    ]:
        pdf.cell(0, 6, _safe(line), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    # Summary
    pdf.set_font(font_name, "", 12)
    pdf.cell(0, 8, _safe("Сводка"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(font_name, "", 10)
    summary_lines = [
        f"Всего требований: {report.total}",
        f"Соответствует: {report.match_count}",
        f"Частично: {report.partial_count}",
        f"Не соответствует: {report.mismatch_count}",
        f"Требует уточнения: {report.clarification_count}",
        f"Баллы: {report.score} / {report.max_score}",
        f"Общее соответствие: {pct}%",
    ]
    for line in summary_lines:
        pdf.cell(0, 6, _safe(line), new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    coverage_rows = _extraction_coverage_rows(report)
    if coverage_rows:
        pdf.set_font(font_name, "", 12)
        pdf.cell(0, 8, _safe("Покрытие извлечения"), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_name, "", 9)
        for warning in _extraction_warning_lines(report):
            pdf.multi_cell(0, 5, _safe(f"Внимание: {warning}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        for row in coverage_rows[:20]:
            pdf.multi_cell(
                0,
                5,
                _safe(f"{row.get('Файл', '')}: {row.get('Показатель', '')} — {_cut(row.get('Значение', ''), 220)}"),
                new_x=XPos.LMARGIN,
                new_y=YPos.NEXT,
            )
        signal_rows = _key_signal_rows(report)
        if signal_rows:
            pdf.set_font(font_name, "", 8)
            for row in signal_rows[:30]:
                pdf.multi_cell(
                    0,
                    4,
                    _safe(
                        f"Сигнал: {row.get('Сигнал', '')}; критичный: {row.get('Критичный', '')}; "
                        f"в извлечении: {row.get('В извлечении', '')}"
                    ),
                    new_x=XPos.LMARGIN,
                    new_y=YPos.NEXT,
                )
        pdf.ln(3)

    if report.summary:
        pdf.set_font(font_name, "", 11)
        pdf.cell(0, 8, _safe("Резюме"), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_name, "", 10)
        pdf.multi_cell(0, 6, _safe(report.summary), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(5)

    refs = _collect_reference_map(report)
    matrix_rows = _platform_matrix_rows(report, refs)
    platform_names = _platform_names(report)
    if matrix_rows and platform_names:
        pdf.add_page(orientation="L")
        pdf.set_font(font_name, "", 13)
        pdf.cell(0, 8, _safe("Матрица соответствия по платформам"), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_name, "", 7)
        headers = ["Пункт ТЗ", "Требование"] + platform_names
        usable_width = pdf.w - pdf.l_margin - pdf.r_margin
        point_width = 25
        requirement_width = min(135, max(80, usable_width * 0.42))
        platform_width = max(18, (usable_width - point_width - requirement_width) / len(platform_names))
        widths = [point_width, requirement_width] + [platform_width] * len(platform_names)
        for width, header in zip(widths, headers):
            pdf.cell(width, 7, _cut(header, 24), border=1)
        pdf.ln(7)
        for row in matrix_rows:
            if pdf.get_y() > 185:
                pdf.add_page(orientation="L")
                pdf.set_font(font_name, "", 7)
                for width, header in zip(widths, headers):
                    pdf.cell(width, 7, _cut(header, 24), border=1)
                pdf.ln(7)
            values = [row.get("Пункт ТЗ", ""), row.get("Требование", "")] + [
                row.get(platform_name, "") for platform_name in platform_names
            ]
            limits = [18, 105] + [max(10, int(platform_width / 2.2))] * len(platform_names)
            for width, value, limit in zip(widths, values, limits):
                pdf.cell(width, 7, _cut(value, limit), border=1)
            pdf.ln(7)
        pdf.add_page()

    totals = _platform_totals(report)
    if totals:
        pdf.set_font(font_name, "", 12)
        pdf.cell(0, 8, _safe("Итого по платформам / услугам"), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_name, "", 9)
        for platform_name, (match_count, partial_count, mismatch_count, clarification_count) in totals.items():
            pdf.multi_cell(
                0,
                5,
                _safe(f"{platform_name}: + {match_count}, ± {partial_count}, - {mismatch_count}, ? {clarification_count}"),
                new_x=XPos.LMARGIN,
                new_y=YPos.NEXT,
            )
        pdf.ln(3)

    external_items = [v for v in report.verdicts if v.requires_external_service]
    if external_items:
        pdf.set_font(font_name, "", 12)
        pdf.cell(0, 8, _safe("Внешние услуги / подрядчики"), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_name, "", 9)
        for v in external_items[:20]:
            pdf.multi_cell(
                0,
                5,
                _safe(f"{v.section or f'#{v.requirement_id}'}: {_req_text(v, 220)} {v.external_service_notes or ''}"),
                new_x=XPos.LMARGIN,
                new_y=YPos.NEXT,
            )
        pdf.ln(3)

    suspicious_items = _suspicious_items(report)
    if suspicious_items:
        pdf.set_font(font_name, "", 12)
        pdf.cell(0, 8, _safe("Сомнительные места"), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_name, "", 9)
        for item in suspicious_items[:25]:
            section = item.get("section") or f"#{item.get('requirement_id')}"
            reasons = "; ".join(item.get("reasons", []) or [])
            verdict_label = VERDICT_LABELS.get(item.get("verdict", ""), item.get("verdict", ""))
            pdf.multi_cell(
                0,
                5,
                _safe(
                    f"{section}: {verdict_label}, уверенность {float(item.get('confidence') or 0):.0%}. "
                    f"Причины: {_cut(reasons, 220)}"
                ),
                new_x=XPos.LMARGIN,
                new_y=YPos.NEXT,
            )
        pdf.ln(3)

    trace_rows = _trace_rows(report)
    if trace_rows:
        pdf.set_font(font_name, "", 12)
        pdf.cell(0, 8, _safe("Трассировка RAG"), new_x="LMARGIN", new_y="NEXT")
        pdf.set_font(font_name, "", 8)
        for row in trace_rows[:40]:
            pdf.multi_cell(
                0,
                4,
                _safe(
                    f"{row.get('Пункт ТЗ', '')}: профиль={row.get('Профиль', '')}; "
                    f"платформа={row.get('Платформа', '')}; score={row.get('Score', '')}; "
                    f"источник={_cut(row.get('Источник', ''), 130)}; "
                    f"причины={_cut(row.get('Причины выбора', ''), 170)}"
                ),
                new_x=XPos.LMARGIN,
                new_y=YPos.NEXT,
            )
        pdf.ln(3)

    pdf.set_font(font_name, "", 12)
    pdf.cell(0, 8, _safe("Что проверить в первую очередь"), new_x="LMARGIN", new_y="NEXT")
    pdf.set_font(font_name, "", 10)
    pdf.multi_cell(0, 6, _safe(_decision_summary(report)), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
    pdf.ln(3)

    sections = [
        ("Несоответствия", sorted([v for v in report.verdicts if v.verdict == "mismatch"], key=_priority_sort_key), "Причина"),
        ("Требуют уточнения", sorted([v for v in report.verdicts if v.verdict == "needs_clarification"], key=_priority_sort_key), "Комментарий"),
        ("Частичное соответствие", sorted([v for v in report.verdicts if v.verdict == "partial"], key=_priority_sort_key), "Обоснование"),
        ("Подтверждённые важные соответствия", _top_key_matches(report), "Почему соответствует"),
    ]

    for title, items, reason_label in sections:
        if not items:
            continue
        pdf.set_font(font_name, "", 13)
        pdf.cell(0, 8, _safe(title), new_x="LMARGIN", new_y="NEXT")
        pdf.ln(2)

        for v in items:
            cat = CATEGORY_LABELS.get(v.category, v.category)
            pdf.set_font(font_name, "", 11)
            pdf.multi_cell(0, 6, _safe(f"{_section_label(v)} ({cat})"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.set_font(font_name, "", 9)
            pdf.multi_cell(0, 5, _safe(f"Пункт / раздел в ТЗ: {v.section or f'#{v.requirement_id}'}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.multi_cell(0, 5, _safe(f"Требование: {_req_text(v, 300)}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            if v.reasoning:
                pdf.multi_cell(0, 5, _safe(f"{reason_label}: {v.reasoning}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            if v.recommendation and v.verdict != "match":
                pdf.multi_cell(0, 5, _safe(f"Рекомендация: {v.recommendation}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            if v.source_urls:
                pdf.multi_cell(0, 5, _safe(f"Документация: {', '.join(v.source_urls[:5])}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(3)

    # Процедурные пункты закупки — вне технической оценки Cloud.ru.
    procedural = sorted(
        [v for v in report.verdicts if (v.category or "").lower() == "procedural"
         or v.verdict == "out_of_scope"],
        key=lambda v: v.requirement_id,
    )
    if procedural:
        pdf.set_font(font_name, "", 13)
        pdf.cell(0, 8, _safe("Процедурные пункты закупки (вне технической оценки)"), new_x="LMARGIN", new_y=YPos.NEXT)
        pdf.set_font(font_name, "", 10)
        pdf.multi_cell(
            0, 6,
            _safe(
                f"Извлечено {len(procedural)} пунктов, относящихся к коммерческо-правовой "
                "обвязке тендера (ОКПД, начальная максимальная цена, обеспечение заявки, "
                "антикоррупция, идентификация участников закупки, реквизиты сторон). Они не "
                "оценивают технические возможности Cloud.ru и исключены из знаменателя процента "
                "соответствия. Передайте их в коммерческую/правовую команду Cloud.ru при "
                "подготовке КП."
            ),
            new_x=XPos.LMARGIN, new_y=YPos.NEXT,
        )
        pdf.ln(2)
        pdf.set_font(font_name, "", 9)
        for v in procedural:
            pdf.multi_cell(
                0, 5,
                _safe(f"{v.section or f'#{v.requirement_id}'}: {_req_text(v, 300)}"),
                new_x=XPos.LMARGIN, new_y=YPos.NEXT,
            )
        pdf.ln(3)

    if refs:
        pdf.set_font(font_name, "", 12)
        pdf.cell(0, 8, _safe("Сноски RAG"), new_x="LMARGIN", new_y=YPos.NEXT)
        pdf.set_font(font_name, "", 8)
        for key, index in sorted(refs.items(), key=lambda item: item[1]):
            pdf.multi_cell(0, 5, _safe(f"[{index}] {key}"), new_x=XPos.LMARGIN, new_y=YPos.NEXT)

    pdf.output(str(path))
    logger.info("Saved PDF report: %s", path)
    return path
